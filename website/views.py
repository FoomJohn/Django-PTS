from django.db import connection, transaction
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.db.models import Count, Avg
from .forms import SignUpForm, AddRecordForm, ScoreForm
from django.contrib.auth.models import User
from .models import Candidate, ScoreEverything, Status, ScoreCard
from .filters import ScoreFilter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def home(request):
    candidates = Candidate.objects.all()
    current_user = request.user
    
    # Initialize candidate_ids as an empty list
    candidate_ids = []

    if current_user.is_authenticated:
        if request.user.is_superuser:
            is_superuser = True
        else:
            is_superuser = False
        # Get all candidate IDs for which the current user has a corresponding entry in the Status model
        candidate_ids = Status.objects.filter(judge=current_user).values_list('candidate_id', flat=True)

    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        # Authentication part
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, "Welcome! You are logged in! :D")
            return redirect('home')
        else:
            messages.error(request, "Invalid username or password.")

    return render(request, 'home.html', {'candidates': candidates, 'candidate_ids': candidate_ids, 'is_superuser': is_superuser})
    
def logout_user(request):
    logout(request)
    messages.success(request, "aight, you logged out, bai!!")
    return redirect('home')


def register_user(request):
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            form.save()
            #autheticate, kinda like autologin after register
            username = form.cleaned_data['username']
            password = form.cleaned_data['password1']
            user = authenticate(username=username, password=password)
            login(request, user)
            messages.success(request, "You is now the registered hehee~")
            return redirect('home')

    else:
        form = SignUpForm()
        return render(request, 'register.html', {'form': form})
    
    return render(request, 'register.html', {'form': form})
    

def candidate_record(request, pk):
    if request.user.is_authenticated:
        #look up the record pk
        current_record = get_object_or_404(Candidate, id=pk)
        candidate_record = Candidate.objects.get(id=pk)
        return render(request, 'record.html', {'candidate_record': candidate_record, 'current_record': current_record})
    else:
        messages.success(request, "You must be logged in to do that")
        return redirect('home')
        

def delete_candidate(request, pk):
    if request.user.is_authenticated:
        delete_it = Candidate.objects.get(id=pk)
        delete_it.delete()
        messages.success(request, "Candidate Deleted")
        return redirect('home')
    else:
        messages.success(request, "You must be logged in to do that")
        return redirect('home')
    


def add_candidate(request):
    
    if request.user.is_authenticated:
        if request.method == "POST":
            form = AddRecordForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                messages.success(request, "Record Added")
                return redirect('home')
        else:
            form = AddRecordForm()
        return render(request, 'add_record.html', {'form': form, })
    else:
        messages.error(request, "You must be logged in to add stuff")
        return redirect('home')


def update_candidate(request, pk):
    if request.user.is_authenticated:
        candidate_record = get_object_or_404(Candidate, id=pk)
        current_record = Candidate.objects.get(id=pk)
        if request.method == "POST":
            form = AddRecordForm(request.POST or None, request.FILES, instance=current_record)
            if form.is_valid():
                form.save()
                messages.success(request, "Record Updated")
                return redirect('home')
        else:
            form = AddRecordForm()
        return render(request, 'update_record.html', {'form':form, 'candidate_record': candidate_record, 'current_record': current_record})
    else:
        messages.success(request, "You must be logged in to add stuff")
        return redirect('home')
#notes
#if person is logging in, they are POST -ing, otherwish they are GET -ing. the request. it's like a bounty in an adveture guild.
    


def score_candidate(request, pk):
    current_record = get_object_or_404(Candidate, id=pk)
    judge = request.user

    # Check if the judge has already scored this candidate
    score_exists = ScoreEverything.objects.filter(candidate=current_record, judge=judge).exists()

    if request.method == "POST":
        form = ScoreForm(request.POST, request.FILES)
        if form.is_valid() and not score_exists:
            instance = form.save(commit=False)
            instance.judge = request.user
            instance.candidate = current_record
            Status.objects.create(judge=request.user, candidate=current_record)
            instance.save()
            
            messages.success(request, "Judged")
            return redirect('home')
    else:
        form = ScoreForm()

    return render(request, 'score_candidate.html', {'form': form, 'current_record': current_record, 'score_exists': score_exists})

def reset_auto_increment(table_name):
    with connection.cursor() as cursor:
        cursor.execute(f"ALTER TABLE {table_name} AUTO_INCREMENT = 1;")


def delete_all_candidate(request):
    if request.user.is_authenticated:
        # Step 1: Delete all candidate objects
        Candidate.objects.all().delete()

        # Step 2: Reset the primary key sequence
        try:
            # Reset auto-increment for website_candidate table
            reset_auto_increment("website_candidate")


            # Commit the transaction
            transaction.commit()

        except Exception as e:
            # Rollback the transaction if any error occurs
            transaction.rollback()
            raise e
        messages.success(request, "All Candidates Removed")
        return redirect('home')  

    else:
        messages.success(request, "You must be logged in to add stuff")
        return redirect('home')
    
    

def scoreeverything_view(request):
    # Initialize variables
    filtered_score_everythings = None

    if request.method == 'POST':
        # Get the selected judge from the form
        selected_judge = request.POST.get('judge')
        if selected_judge:
            # Filter score_everythings based on the selected judge
            filtered_score_everythings = ScoreEverything.objects.filter(judge=selected_judge)

    # Fetch all score entries from the database if no filter is applied
    if filtered_score_everythings is None:
        filtered_score_everythings = ScoreEverything.objects.all()

    return render(request, 'scoreeverything.html', {'filtered_score_everythings': filtered_score_everythings})







def tabulation(request):
    
    scoreeverythings = ScoreEverything.objects.all().order_by('candidate', 'judge')

    scorefilter = ScoreFilter(request.GET, queryset=scoreeverythings)
    scoreeverythings = scorefilter.qs

    if request.user.is_authenticated:
        return render(request, 'tabulation.html', {'scoreeverythings':scoreeverythings, 'scorefilter':scorefilter, })
    else:
        messages.success(request, "noo")
        return redirect('home')
    
def tabulation_production_number(request):
    
    scoreeverythings = ScoreEverything.objects.all().order_by('candidate', 'judge')

    scorefilter = ScoreFilter(request.GET, queryset=scoreeverythings)
    scoreeverythings = scorefilter.qs


    if request.user.is_authenticated:
        return render(request, 'tabulation_production_number.html', {'scoreeverythings':scoreeverythings, 'scorefilter':scorefilter,})
    else:
        messages.success(request, "noo")
        return redirect('home')      

def tabulation_swimsuit(request):
    
    scoreeverythings = ScoreEverything.objects.all().order_by('candidate', 'judge')

    scorefilter = ScoreFilter(request.GET, queryset=scoreeverythings)
    scoreeverythings = scorefilter.qs


    if request.user.is_authenticated:
        return render(request, 'tabulation_swimsuit.html', {'scoreeverythings':scoreeverythings, 'scorefilter':scorefilter,})
    else:
        messages.success(request, "noo")
        return redirect('home')   
    
def tabulation_evening_gown(request):
    
    scoreeverythings = ScoreEverything.objects.all().order_by('candidate', 'judge')

    scorefilter = ScoreFilter(request.GET, queryset=scoreeverythings)
    scoreeverythings = scorefilter.qs


    if request.user.is_authenticated:
        return render(request, 'tabulation_evening_gown.html', {'scoreeverythings':scoreeverythings, 'scorefilter':scorefilter,})
    else:
        messages.success(request, "noo")
        return redirect('home')   
    
def tabulation_q_and_a(request):
    
    scoreeverythings = ScoreEverything.objects.all().order_by('candidate', 'judge')

    scorefilter = ScoreFilter(request.GET, queryset=scoreeverythings)
    scoreeverythings = scorefilter.qs


    if request.user.is_authenticated:
        return render(request, 'tabulation_q_and_a.html', {'scoreeverythings':scoreeverythings, 'scorefilter':scorefilter,})
    else:
        messages.success(request, "noo")
        return redirect('home')   

def tabulation_calculate(request):
    
    scoreeverythings = ScoreEverything.objects.all().order_by('candidate', 'judge')
    scorecard = ScoreCard.objects.all().order_by('ranking')

    scorefilter = ScoreFilter(request.GET, queryset=scoreeverythings)
    scoreeverythings = scorefilter.qs

    unique_judges_count = Status.objects.values('judge').distinct().count()
    print("Unique Judges Count:", unique_judges_count)
    
    # Get the count of candidates
    max_candidate_count = Candidate.objects.count()
    print("Max Candidate Count:", max_candidate_count)

    # Get the total count of entries in the Status table
    total_status_entries = Status.objects.count()
    print("Total Status Entries:", total_status_entries)

    # Calculate the expected count of entries if each judge appears for every candidate
    expected_entries_count = unique_judges_count * max_candidate_count
    print("Expected Entries Count:", expected_entries_count)

    # Check if the actual count of entries matches the expected count
    judge_done = total_status_entries == expected_entries_count
    print("Judge Done:", judge_done)

    candidates = Candidate.objects.all()

    # Iterate through each candidate
    

    if request.user.is_authenticated:
        if request.user.is_superuser:
            is_superuser = True
        else:
            is_superuser = False
        return render(request, 'tabulation_calculate.html', {'scoreeverythings':scoreeverythings, 'scorefilter':scorefilter, 'judge_done':judge_done, 'is_superuser': is_superuser, 'scorecard': scorecard})
    else:
        messages.success(request, "noo")
        return redirect('home')


def calculate_scorecard(request):
    # Get all candidates
    candidates = Candidate.objects.all()

    # Iterate through each candidate
    for candidate in candidates:
        # Get all ScoreEverything instances for the current candidate
        scores = ScoreEverything.objects.filter(candidate=candidate)

        # Calculate the average of pn_total for the candidate
        pn_avg = scores.aggregate(pn_avg=Avg('pn_total'))['pn_avg']

        # Calculate the average of sw_total for the candidate
        sw_avg = scores.aggregate(sw_avg=Avg('sw_total'))['sw_avg']

        # Calculate the average of eg_total for the candidate
        eg_avg = scores.aggregate(eg_avg=Avg('eg_total'))['eg_avg']

        # Calculate the average of fq_total for the candidate
        fq_avg = scores.aggregate(fq_avg=Avg('fq_total'))['fq_avg']

        # Calculate the average of t_avg for the candidate
        t_avg = scores.aggregate(t_avg=Avg('t_avg'))['t_avg']

        # Create or update the ScoreCard entry for the candidate
        score_card, created = ScoreCard.objects.get_or_create(candidate=candidate)
        score_card.pn_all_total = pn_avg
        score_card.sw_all_total = sw_avg
        score_card.eg_all_total = eg_avg
        score_card.fq_all_total = fq_avg
        score_card.t_all_avg = t_avg
        score_card.save()

    # Get all ScoreCard objects and order them by total average score
    scorecards = ScoreCard.objects.all().order_by('-t_all_avg')

    # Assign ranks based on the sorted order
    
    rank = 1
    prev_score = None
    prev_rank = None
    for scorecard in scorecards:
        if prev_score is not None and scorecard.t_all_avg == prev_score:
            # If the current score is the same as the previous score, assign the same rank
            scorecard.ranking = prev_rank
        else:
            # Otherwise, assign a new rank
            scorecard.ranking = rank
            prev_rank = rank
        scorecard.save()
        prev_score = scorecard.t_all_avg
        rank += 1
        

    messages.success(request, "Calculated Averages and Rankings!")
    return redirect('tabulation_calculate')

"""
def generate_scorecard_document(request):
    # Get ScoreCard objects ordered by ranking
    scorecards = ScoreCard.objects.order_by('ranking')

    # Create a new Word document
    doc = Document()
    
    # Add a title to the document
    doc.add_heading('ScoreCard', level=1)

    # Define the number of columns in the table
    num_columns = 7

    # Add a table to the document
    table = doc.add_table(rows=1, cols=num_columns)
    table.style = 'Table Grid'

    # Add column headers to the table
    headings = ['Ranking', 'Candidate', 'PN Total', 'SW Total', 'EG Total', 'FQ Total', 'Total Avg']
    header_cells = table.rows[0].cells
    if len(headings) == num_columns:
        print("len(headings) == num_columns is True")
        for i, heading in enumerate(headings):
            header_cells[i].text = heading
    else:
        print("len(headings) == num_columns is False")
        # Handle the case where the number of headings doesn't match the number of columns
        return redirect('tabulation_calculate')  # Redirect to a suitable page or handle the error

    # Add ScoreCard data to the table rows
    for scorecard in scorecards:
        row = table.add_row().cells
        row[0].text = str(scorecard.ranking)
        row[1].text = f"{scorecard.candidate.first_name} {scorecard.candidate.last_name}"
        row[2].text = str(scorecard.pn_all_total)
        row[3].text = str(scorecard.sw_all_total)
        row[4].text = str(scorecard.eg_all_total)
        row[5].text = str(scorecard.fq_all_total)
        row[6].text = str(scorecard.t_all_avg)

    # Save the document to a temporary file
    doc_file_path = 'website/documents/scorecard_document.docx'
    doc.save(doc_file_path)

    # Redirect the user back to the tabulation calculate page
    return redirect('tabulation_calculate')"""


def generate_scorecard_document(request):

    try:
        # Get ScoreCard objects ordered by ranking
        scorecards = ScoreCard.objects.order_by('ranking')

        # Create a new Word document
        doc = Document()
        
        # Add a title to the document
        title = doc.add_heading('ScoreCard', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.style.font.bold = True
        title.style.font.size = Pt(24)

        # Add a table to the document
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        # Add column headers to the table
        headings = ['Ranking', 'Candidate', 'PN Total', 'SW Total', 'EG Total', 'FQ Total', 'Total Avg']
        header_cells = table.rows[0].cells
        for i, heading in enumerate(headings):
            header_cells[i].text = heading

        # Add ScoreCard data to the table rows
        for scorecard in scorecards:
            row = table.add_row().cells
            row[0].text = str(scorecard.ranking)
            row[1].text = f"{scorecard.candidate.first_name} {scorecard.candidate.last_name}"
            row[2].text = str(scorecard.pn_all_total)
            row[3].text = str(scorecard.sw_all_total)
            row[4].text = str(scorecard.eg_all_total)
            row[5].text = str(scorecard.fq_all_total)
            row[6].text = str(scorecard.t_all_avg)

        # Save the document to the specified file path
        doc_file_path = 'website/documents/scorecard_document.docx'
        doc.save(doc_file_path)

        messages.success(request, "Printed Scores!")
        # Redirect the user back to the tabulation calculate page
        return redirect('tabulation_calculate')

    except PermissionError as e:
            # Display a message if the document is open and cannot be saved
            messages.error(request, "Can't print scores when Document is open")
            return redirect('tabulation_calculate')